#export TESSDATA_PREFIX=/usr/share/tesseract-ocr/5/tessdata
import pytesseract
import os
from PIL import Image
from pdf2image import convert_from_path
from docx import Document


def extract_bangla_text_with_tesseract(image_path, lang='ben'):
    image = Image.open(image_path)
    text = pytesseract.image_to_string(image, lang=lang)
    return text



fi = 0
doc = Document()
for file in os.listdir("output"):
            fi += 1
            output_image_path = f'output/{fi}.png'  
            
            bangla_text = extract_bangla_text_with_tesseract(output_image_path)
            
            #print(bangla_text + "\n\n--------------------------------------------\n\n")
            print('Extracting:  ' + f'{int(fi/31*100)} %')   
            
            doc.add_paragraph("Page: " + str(fi) + "\n----------------------------------------------------------\n") 
            doc.add_paragraph(bangla_text + "\n\n-----------------------------------------------------------\n\n")    
            #result = pytesseract.image_to_data(image, lang='ben', output_type=pytesseract.Output.DICT)
            result = pytesseract.image_to_data(output_image_path, lang='ben', output_type=pytesseract.Output.DICT)
            # here I could write 'eng+ben' as well to extract bangla+english
            for i, word_text in enumerate(result['text']):
                confidence = int(result['conf'][i])
                doc.add_paragraph(f"Word: {word_text}, Confidence: {confidence}")
                #print(f"Word: {word_text}, Confidence: {confidence}")
            doc.add_paragraph("End  Page-" + str(fi) + "\n----------------------------------------------------------\n")
doc.save('ConfidenceTest.docx')

