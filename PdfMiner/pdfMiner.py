#import PyPDF2
from pdfminer.high_level import extract_text
import os
from docx import Document

# def extract_text_using_pypdf2(pdf_path):
#     text = ''
#     with open(pdf_path, 'rb') as pdf_file:
#         reader = PyPDF2.PdfReader(pdf_file)
#         for page_num in range(len(reader.pages)):
#             page = reader.pages[page_num]
#             text += page.extract_text()
#     return text

def extract_text_using_pdfminer(pdf_path):
    text = extract_text(pdf_path)
    return text


input_directory = "../../PDF/tryPDF"

for filename in os.listdir(input_directory):
    if filename.endswith(".pdf"):
        pdf_path = os.path.join(input_directory, filename)
        #pdf_path = 'path_to_your_pdf_file.pdf'

        # Using PyPDF2
        #text_pypdf2 = extract_text_using_pypdf2(pdf_path)
        #print("Extracted text using PyPDF2:")
        #print(text_pypdf2)

        # Using pdfminer
        text_pdfminer = extract_text_using_pdfminer(pdf_path)
        print("Extracted text using pdfminer:")
        print(text_pdfminer)

        # Create a new Document object
        doc = Document()

        # Add content to the document
        doc.add_paragraph("")
        doc.add_paragraph(text_pdfminer)
        doc.add_paragraph("\n\n\n\n")
        # Save the document
        file_path = "PdfMiner.docx"  # Replace with your desired file path
        doc.save(file_path)
