import pdfplumber
import os

from docx import Document


def extract_bangla_text(pdf_path):
    text = ''
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            text += page_text
    return text

input_directory = "../../PDF/tryPDF"

# Replace "pdf_path" with the actual path to your PDF file
for filename in os.listdir(input_directory):
    if filename.endswith(".pdf"):
        pdf_path = os.path.join(input_directory, filename)

        bangla_text = extract_bangla_text(pdf_path)
        print("Extracted Bangla text:")
        print(bangla_text)




        # Create a new Document object
        doc = Document()

        # Add content to the document
        doc.add_paragraph("")
        doc.add_paragraph(bangla_text)
        doc.add_paragraph("\n\n\n\n")
        # Save the document
        file_path = "PdfPlumber.docx"  # Replace with your desired file path
        doc.save(file_path)
