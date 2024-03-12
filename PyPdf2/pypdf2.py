import PyPDF2
#import re
import os
from docx import Document

def extract_text_from_pdf(pdf_path):
    text = ""
    with open(pdf_path, "rb") as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        for page in pdf_reader.pages:
            text += page.extract_text()

    # Remove page numbers, headers, footers
    #text = re.sub(r'\b\d+\b', '', text)  # Remove page numbers
    #text = re.sub(r'\n\s*\n', '\n', text)  # Remove empty lines
    #text = re.sub(r'\[[^\]]+\]', '', text)  # Remove header
    #text = re.sub(r'\[[^\]]+\]$', '', text, flags=re.MULTILINE)  # Remove footer
    #text = re.sub(r'\b\w+\b', "", text, flag=re.IGNORECASE) # Remove English Word
    # Add more regex patterns to remove specific headers/footers

    return text

input_directory = "../../PDF/tryPDF"

# Replace "pdf_path" with the actual path to your PDF file
for filename in os.listdir(input_directory):
    if filename.endswith(".pdf"):
        pdf_path = os.path.join(input_directory, filename)
        extracted_text = extract_text_from_pdf(pdf_path)
        print(extracted_text)
        # Create a new Document object
        doc = Document()

        # Add content to the document
        doc.add_paragraph("")
        doc.add_paragraph(extracted_text)
        doc.add_paragraph("\n\n\n\n")
        # Save the document
        file_path = "PyPDF2.docx"  # Replace with your desired file path
        doc.save(file_path)
        
# Open the file in append mode
"""
file_path = "text.txt"  # Replace with your desired file path
with open(file_path, "a") as file:
    # Write new content to the file
    file.write(extracted_text)
    file.write("\n\n\n\n\n\n")
"""

